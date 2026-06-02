import os
import asyncio
import aiofiles
import aiohttp
import base64
import httpx
import json
import threading
import urllib.request
from io import StringIO, BytesIO
import logging
from urllib.error import HTTPError, URLError
from typing import Any, Dict, List
from pathlib import Path

from PIL import Image

from docx import Document
from docx.shared import Inches, RGBColor, Mm, Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK

from .local_image_service import local_image_service
from .s3_image_service import s3_image_service
from .repository import Report

logger = logging.getLogger(__name__)

class DocxRenderer:

    def __init__(self):
        self.doc = Document()
        self.logoname = "thermoelectrica_logo.png"

    @property
    def section(self):
        # доступ к первой секции
        section = self.doc.sections[0]  # доступ к первому разделу
        # высота листа в сантиметрах
        section.page_height = Cm(29.7)
        # ширина листа в сантиметрах
        section.page_width = Cm(21.0)
        # левое поле в миллиметрах
        section.left_margin = Mm(20.4)
        # правое поле в миллиметрах
        section.right_margin = Mm(10)
        # верхнее поле в миллиметрах
        section.top_margin = Mm(15)
        # нижнее поле в миллиметрах
        section.bottom_margin = Mm(10)
        # отступ от верхнего края страницы до 
        # нижнего края нижнего колонтитула
        section.header_distance = Mm(10)
        # отступ от нижнего края страницы до 
        # нижнего края нижнего колонтитула
        section.footer_distance = Mm(10)
        return section


    async def fetch_async_data(self, client, item):
        try:
            file_name = f"./data_images/{item['name']}"
            if True: #not os.path.isfile(file_name):
            #if not file_path.exists():
                total = 0
                async with client.stream("GET", item["url"], follow_redirects=True) as response:
                    response.raise_for_status()

                    file_size = int(response.headers.get("content-length", 0))
                    print(f"Размер файла: {file_size / 1024 / 1024:.1f} МБ")

                    with open(file_name, "wb") as file:
                        async for chunk in response.aiter_bytes(chunk_size=8192):
                            file.write(chunk)
                            total += len(chunk)

                            if file_size:
                                pct = total / file_size * 100
                                print(f"\rЗагружено: {pct:.1f}%", end="", flush=True)

                    # Изменение размера изображения
                    with Image.open(file_name) as img:
                        width, height = img.size
                        new_width = 400
                        new_height = int(height * (new_width / width))
                        resized_img = img.resize((new_width, new_height), Image.BICUBIC)
                        resized_img.save(file_name)  # Сохранение изменённого изображения
                return True

        except Exception as exc:
            #print(f"Ошибка: {exc}")
            logger.info(f"Error: {exc}")
            return f"Error: {exc}"

    async def fetch_images(self, query_results: Dict[str, List[Dict[str, Any]]]):
        urls_collection = []
        for row in query_results["data"]:
            for image in ["visual_image_id", "thermal_image_id"]:
                if row.get(image):
                    url = s3_image_service.image_url(row.get(image))
                    urls_collection.append({"name": row.get(image), "url": url})

        async with httpx.AsyncClient() as client:
            #print(f"URLS COLLECTION: {urls_collection}")
            tasks = [self.fetch_async_data(client, item) for item in urls_collection]
            results = await asyncio.gather(*tasks, return_exceptions=True)
            for i, result in enumerate(results):
                if result:
                    print(f"Результат {i+1} получен: {result}")


    async def render(
        self,
        report: Report,
        params: Dict[str, Any],
        query_results: Dict[str, List[Dict[str, Any]]],
    ) -> bytes:
        
        self.query_results = query_results
        await asyncio.create_task(self.fetch_images(self.query_results))
        return await self.render_docx(report, params)
      

    async def render_docx(self, report: Report, params: Dict[str, Any]) -> bytes:
     
        print(f"RENDER DOCX...")
        print(f"RENDER DOCX, REPORT: {report}")
        print(f"RENDER DOCX, PARAMETERS: {params}")
        print(f"RENDER DOCX, QUERY RESULTS: {self.query_results}")

        logo = local_image_service.get_image_data_uri(
            self.logoname, report.path
        )
        # доступ к первой секции
        #section = self.doc.sections[0]  # доступ к первому разделу
  
        # Межстрочный интервал
        style = self.doc.styles['Normal']
        style.paragraph_format.line_spacing = 1.1

        # доступ к верхнему колонтитулу
        header = self.section.header.paragraphs[0]

        # добавляем логотип слева в хедер
        if logo:
            logo_run = header.add_run()
            logo_run.add_picture(BytesIO(base64.b64decode(logo.split(",")[1])), width=Inches(1.5))

        # Добавление текста справа в хедер
        text_run = header.add_run()
        date_format = params["period_end"].strftime("«%d» %B %Y г.") if params["period_end"] else ""
        text_run.text = f"\t\tПРОТОКОЛ № {params['protocol_number']} { params['plant_name'] } от { date_format }"
        text_run.font.size = Pt(8)
        #text_run.bold = True
        #text_run.alignment = WD_ALIGN_PARAGRAPH.RIGHT

         # доступ к нижнему колонтитулу
        footer = self.section.footer.paragraphs[0]

        # добавляем нижний колонтитул
        footer.add_run("* согласно РД 34.45-51.300-97 и типовой инструкции по применению термоиндикаторов\t\t")
        footer.style.font.size = Pt(8)
        # выравниваем колонтитул по правому краю
        #footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # добавляем заголовок документа 
        title = self.doc.add_heading(
            (
                f"Протокол теплового контроля контактов и контактных соединений" 
                f" с применением термоиндикаторов на { params['plant_name'] }"
            ), 
            1
        ) #.add_break(WD_BREAK.LINE)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # создаем параграф
        p = self.doc.add_paragraph()

        # вставляем пустую строку
        run = p.add_run()
        run.add_break(WD_BREAK.LINE)

        # добавляем номер протокола и название
        p.add_run(
            (
                f"ПРОТОКОЛ № {params['protocol_number']} { params['plant_name'] } от { date_format }\n"
                f"теплового контроля электрооборудования"
            )
        )
        #run = p.add_run()
        #run.add_break()
        #p.add_run(
            #(
                #f"ПРОТОКОЛ № {params['protocol_number']} { params['plant_name'] } от { date_format }\n"
                #f"теплового контроля электрооборудования"
            #)
        #)
        #run.font.size = Pt(9)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # вставляем пустую строку
        run = p.add_run()
        run.add_break(WD_BREAK.LINE)

        # Таблица 1
        table1 = self.doc.add_table(rows=1, cols=2)
        table1.style = 'Table Grid'

        # Установка размеров ячеек Таблицы 1
        table1.autofit = False # Отключает автоматическое выравнивание ширины
        table1.allow_autofit = False # Гарантирует, что Word не переопределит макет
        table1.columns[0].width = Inches(3.0)
        #table1.rows[0].cells[0].width = Inches(3.0)
        table1.columns[1].width = Inches(4.0)
        #table1.rows[0].cells[1].width = Inches(4.0)
    
        # Текст в ячейке 0,0
        #hdr_cells = table1.rows[0].cells
        table1.cell(0,0).text = "Перечень обследованных электроустановок"

        # Размер шрифта в ячейке 0,0
        paragraph = table1.rows[0].cells[0].paragraphs[0]
        #run = paragraph.runs
        font = paragraph.runs[0].font
        font.size= Pt(7.5)
        font.bold = True

        equipment = list(
            set(
                [ item["equipment_type_name"] for item in self.query_results["data"] ]
            )
        )
        equipment_string = ("\n").join(equipment)

        control_points_count, control_tin_count = 0, 0
        if self.query_results["inspection_summary"]:
            control_points_count = sum([item["total_point_count"] for item in self.query_results["inspection_summary"]])
            control_tin_count = sum([item["total_sticker_count"] for item in self.query_results["inspection_summary"]])
        else:
            control_points_count = len(self.query_results["data"])
            control_tin_count = sum(
                [
                    1 for item in self.query_results["data"]
                    if item["is_sticker_present"] is True
                ]
            )

        control_points_string = f"Количество контрольных точек: {control_points_count}"
        control_tin_string = f"Количество ТИН: {control_tin_count}"

        table1.cell(0,1).text = (
            equipment_string + "\n" 
            + control_points_string + "\n" 
            + control_tin_string + "\n"
        )

        # Размер шрифта в ячейке 0,1
        paragraph = table1.rows[0].cells[1].paragraphs[0]
        font = paragraph.runs[0].font
        font.size= Pt(7.5)

        # переход на следующую страницу
        p = self.doc.add_paragraph()
        run = p.add_run()
        run.add_break(WD_BREAK.LINE)

        ### Таблица 3-1 (Заголовки 1) ###
        tab3_columns = 4
        table3 = self.doc.add_table(rows=1, cols=4)
        table3.style = 'Table Grid'

        # Определение цвета фона и текста
        #header_bg = RGBColor(20, 50, 100)  # Тёмно-синий
        #header_text_color = RGBColor(255, 255, 255)  # Белый

        # Установка размеров ячеек Таблицы 3
        table3.autofit = False # Отключает автоматическое выравнивание ширины
        table3.allow_autofit = False # Гарантирует, что Word не переопределит макет
        table3.columns[0].width = Inches(0.5)
        table3.columns[1].width = Inches(1.5)
        table3.columns[2].width = Inches(1.5)
        table3.columns[3].width = Inches(3.5)
        #table3.rows[0].cells[0].width = Inches(0.5)
        #table3.columns[1].width = Inches(4.0)
        #table3.rows[0].cells[1].width = Inches(4.0)

        # Устанавливаем заголовки таблицы
        hdr_cells = table3.rows[0].cells
        hdr_cells[0].text = "№"
        hdr_cells[1].text = "Диспетчерское наименование электрооборудования; узел"
        hdr_cells[2].text = "Фотография термоиндикатора и термограмма"
        hdr_cells[3].text = "Выявленный дефект"

        # Стили шрифта в заголовке Таблицы 3
        paragraph = table3.rows[0].cells[0].paragraphs[0]
        #run = paragraph.runs
        font = paragraph.runs[0].font
        font.size= Pt(7.5)
        font.bold = True

        ### Таблица 3-2 (Заголовок 2) ###
        table3_2 = self.doc.add_table(rows=1, cols=1)
        table3_2.style = 'Table Grid'
        #table3.columns[0].width = Inches(7)
        #hdr_cells = table3.rows[0].cells
        #hdr_cells[0].text = "Аварийные дефекты распределительных устройств"

        ### Таблица 3-3 (Ряды данных) ###
        table3 = self.doc.add_table(rows=0, cols=4)
        table3.style = 'Table Grid'

        # Определение цвета фона и текста
        #header_bg = RGBColor(20, 50, 100)  # Тёмно-синий
        #header_text_color = RGBColor(255, 255, 255)  # Белый

        # Установка размеров ячеек Таблицы 3-3
        table3.columns[0].width = Inches(0.5)
        table3.columns[1].width = Inches(2.5)
        table3.columns[2].width = Inches(1.0)
        table3.columns[3].width = Inches(3.0)

        for idx, row in enumerate(self.query_results["data"]):
            group_label = ""
            if row["criticality"] == "CRITICAL" and row["is_panel"] == "MOTOR":
                group_label = "Критические дефекты двигателей"
            elif row["criticality"] == "CRITICAL" and row["is_panel"] == "PANEL":
                group_label = "Критические дефекты распределительных устройств"
            elif row["criticality"] == "EMERGENCY" and row["is_panel"] == "MOTOR":
                group_label = "Аварийные дефекты двигателей"
            elif row["criticality"] == "EMERGENCY" and row["is_panel"] == "PANEL":
                group_label = "Аварийные дефекты распределительных устройств"
            elif row["criticality"] == "DEVELOPING" and row["is_panel"] == "MOTOR":
                group_label = "Развивающиеся дефекты двигателей"
            else:
                group_label = "Развивающиеся дефекты распределительных устройств"
            if group_label:
                table3_2.columns[0].width = Inches(7)
                hdr_cells = table3_2.rows[0].cells
                hdr_cells[0].text = group_label

            row_cells = table3.add_row().cells
            row_cells[0].text = str(idx + 1)

            paragraph = row_cells[1].paragraphs[0]
            
            for image in ["visual_image_id", "thermal_image_id"]:
                if row.get(image):
                    print(f"GET IMAGE: {f'./data_images/{row.get(image)}'}")
                    paragraph.add_run().add_picture(f"./data_images/{row.get(image)}", width=Cm(8))
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    paragraph.add_run("")
            #paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            #row_cells[1].text = ""
            row_cells[2].text = ""
            row_cells[3].text = ""

        self.doc.save('demo_report.docx')

        try:
            # Read the generated DOCX file
            with open('./demo_report.docx', 'rb') as f:
                docx_bytes = f.read()
        #finally:
                # Clean up temporary files
                #Path(output_file_path).unlink(missing_ok=True)
        except Exception as e:
            logger.error(f"Error: {e}")

        logger.info(f"DOCX generated successfully, size: {len(docx_bytes)} bytes")
        return docx_bytes